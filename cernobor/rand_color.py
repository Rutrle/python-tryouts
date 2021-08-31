from random import randint
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Mm
import docx


def create_rand_list(main_length, list_length, max_num):
    """
    creates list of random lists, each unique containing numbers in range from 1 to max_num

    :param main_length: int length of the main list
    :param list_length: int lengths of the random lists
    :param max_num: int
    """
    rand_list = []
    while len(rand_list) < main_length:
        single_list = []

        for i in range(list_length):
            single_list.append(randint(1, max_num),)

        if single_list not in rand_list and single_list.reverse() not in rand_list:

            rand_list.append(single_list)

    return rand_list


def write_to_file(rand_list):
    """
    write color codes into newly created .docx file
    :param rand_list: list of lists
    """

    color_dict = {1: 'FF0000', 2: 'FFCA00',
                  3: '88D0FF', 4: '007800', 5: 'A8A9AD'}

    info_list = ['', '15+ orlénů', '10 orlénů',
                 '6 orlénů', '3 orlény', '1 orlén', '1f']
    info_count = 0

    mydoc = docx.Document()

    for single_list in rand_list:

        if info_count % 6 == 0 or info_count == 0:
            para = mydoc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(5)

            current_run = para.add_run(info_list.pop())
            current_run.font.size = Pt(10)

        info_count = info_count+1

        para = mydoc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)

        for point in single_list:
            if point == 6:
                # n is a full square in Wingdings
                current_run = para.add_run('p')
                current_run.font.name = 'Wingdings'
                current_run.font.size = Pt(7)
            else:
                # n is a full square in Wingdings
                current_run = para.add_run('n')
                current_run.font.name = 'Wingdings'

                current_run.font.color.rgb = RGBColor.from_string(
                    color_dict[point])

                current_run.font.size = Pt(8)

    mydoc.save('color_code.docx')


rand_list = create_rand_list(36, 5, 6)

write_to_file(rand_list)
